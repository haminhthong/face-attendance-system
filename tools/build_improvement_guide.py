from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(__file__).resolve().parents[1] / "HUONG_DAN_CAI_THIEN_CHI_TIET.docx"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
GREEN = "E2F0D9"
YELLOW = "FFF2CC"
RED = "FCE4D6"
WHITE = "FFFFFF"
TEXT = "243447"


def set_cell_fill(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def add_table(doc, headers, rows, widths, status_col=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = 0
    header = table.rows[0]
    set_repeat_table_header(header)
    for index, text in enumerate(headers):
        cell = header.cells[index]
        set_cell_fill(cell, BLUE)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(text)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(WHITE)
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = str(value)
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if status_col == index:
                color = GREEN if str(value).startswith("Đạt") else YELLOW if "phần" in str(value) else RED
                set_cell_fill(cells[index], color)
    set_table_geometry(table, widths)
    doc.add_paragraph()
    return table


def add_bullet(doc, text, level=0):
    paragraph = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    paragraph.add_run(text)


def add_number(doc, text):
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.add_run(text)


def add_callout(doc, title, text, fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_fill(cell, fill)
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(f"{title}: ")
    run.bold = True
    paragraph.add_run(text)
    doc.add_paragraph()


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Title", 28, DARK_BLUE, 0, 10),
        ("Subtitle", 13, "5B6573", 0, 12),
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Bullet 2", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    if "Code" not in styles:
        code = styles.add_style("Code", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code = styles["Code"]
    code.font.name = "Consolas"
    code.font.size = Pt(9)
    code.paragraph_format.left_indent = Inches(0.25)
    code.paragraph_format.space_before = Pt(4)
    code.paragraph_format.space_after = Pt(8)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("Face Attendance System  |  Hướng dẫn cải thiện")


def add_code(doc, text):
    paragraph = doc.add_paragraph(style="Code")
    paragraph.add_run(text)


def build_document():
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("HƯỚNG DẪN CẢI THIỆN CHI TIẾT")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Face Attendance System - từ prototype đến dự án portfolio đáng tin cậy")
    doc.add_paragraph()
    add_callout(
        doc,
        "Mục tiêu tài liệu",
        "Chuyển kết quả rà soát bốn tầng thành kế hoạch kỹ thuật có thứ tự, tiêu chí hoàn thành và bằng chứng cần thu thập. Tài liệu không khẳng định các chỉ số AI chưa được đo.",
    )
    add_table(
        doc,
        ["Thông tin", "Nội dung"],
        [
            ("Phạm vi", "Problem → AI/ML correctness → Software Engineering → Production/Business value"),
            ("Đối tượng", "Chủ dự án, người review GitHub/CV và người tiếp tục phát triển"),
            ("Hiện trạng kiểm chứng", "Mã nguồn module hóa; 12 test đã chạy thành công ở lần kiểm tra gần nhất"),
            ("Nguyên tắc", "Không bịa metric, không đưa dữ liệu sinh trắc học thật lên GitHub, không gọi prototype là production-ready"),
        ],
        [2160, 7200],
    )

    doc.add_heading("Mục lục", level=1)
    for item in (
        "1. Kết luận và nguyên tắc cải thiện",
        "2. Tầng 1 - Problem",
        "3. Tầng 2 - AI/ML correctness",
        "4. Tầng 3 - Software Engineering",
        "5. Tầng 4 - Production và Business value",
        "6. Lộ trình triển khai theo sprint",
        "7. Ma trận kiểm thử và tiêu chí nghiệm thu",
        "8. Hướng dẫn trình bày trên GitHub và CV",
    ):
        add_bullet(doc, item)
    doc.add_page_break()

    doc.add_heading("1. Kết luận và nguyên tắc cải thiện", level=1)
    doc.add_paragraph(
        "Dự án đang giải quyết đúng nhu cầu điểm danh khuôn mặt ở mức prototype. Phần kỹ thuật phần mềm tương đối tốt, nhưng chưa có bộ đánh giá độc lập để chứng minh độ chính xác AI, chưa có load test và chưa đủ kiểm soát bảo mật cho triển khai thật. Vì vậy, hướng cải thiện tốt nhất là tăng bằng chứng kỹ thuật thay vì thêm nhiều thành phần AI không cần thiết."
    )
    add_table(
        doc,
        ["Tầng", "Hiện trạng", "Mục tiêu tiếp theo"],
        [
            ("Problem", "Đạt phần lớn", "Đóng phạm vi và bổ sung quy trình giáo viên xác nhận"),
            ("AI/ML correctness", "Chưa đủ bằng chứng", "Dataset đánh giá riêng, calibration và metric thật"),
            ("Software Engineering", "Khá tốt", "Tăng edge-case test, lint, type-check và clean-install"),
            ("Production/Business", "Prototype", "RBAC, bảo vệ biometric, load test và vận hành"),
        ],
        [1800, 2160, 5400],
    )
    add_callout(doc, "Tuyên bố phù hợp", "Có thể ghi đây là hệ thống điểm danh khuôn mặt dạng prototype có UI, API, database, kiểm tra tương tác và automated tests.", GREEN)
    add_callout(doc, "Không nên tuyên bố", "Không ghi accuracy, production-ready, chống giả mạo an toàn hoặc hỗ trợ 100 camera đồng thời khi chưa có kết quả đo.", RED)

    doc.add_heading("2. Tầng 1 - Problem", level=1)
    doc.add_heading("2.1. Chuẩn hóa bài toán", level=2)
    doc.add_paragraph("Định nghĩa đề xuất:")
    add_callout(doc, "Problem statement", "Hệ thống hỗ trợ giáo viên ghi nhận sự hiện diện của sinh viên trong một buổi học bằng đối chiếu khuôn mặt đã đăng ký, có kiểm tra tương tác cơ bản, ngăn bản ghi trùng và cho phép truy vết kết quả.")
    doc.add_paragraph("Phạm vi nên được ghi rõ:")
    for text in (
        "Dùng cho lớp học nhỏ hoặc vừa; không phải hệ thống kiểm soát an ninh.",
        "Kết quả AI là đề xuất; giáo viên có quyền xác nhận hoặc sửa.",
        "Sinh viên phải đồng ý trước khi lưu embedding khuôn mặt.",
        "Không dùng kết quả tự động làm căn cứ kỷ luật nếu không có xác minh con người.",
    ):
        add_bullet(doc, text)

    doc.add_heading("2.2. Chuẩn hóa Input → Output", level=2)
    add_table(
        doc,
        ["Giai đoạn", "Đầu vào", "Đầu ra/điều kiện"],
        [
            ("Đăng ký", "Ảnh RGB + mã sinh viên + consent", "Đúng 1 khuôn mặt, đạt chất lượng, embedding 128 chiều"),
            ("Nhận diện", "Frame camera + roster + template", "Ứng viên Top-1, distance, margin và lý do từ chối"),
            ("Liveness", "Chuỗi frame và EAR", "Đạt/không đạt challenge, không được xem là PAD hoàn chỉnh"),
            ("Điểm danh", "Identity đã chấp nhận + buổi học", "present/late hoặc mã lỗi nghiệp vụ"),
        ],
        [1500, 3420, 4440],
    )
    doc.add_paragraph("Kết quả nên có cấu trúc rõ ràng, không gọi 1 - distance là xác suất tin cậy:")
    add_code(doc, '{\n  "student_id": "SV001",\n  "decision": "accepted",\n  "distance": 0.42,\n  "liveness_passed": true,\n  "status": "present",\n  "recognized_at": "ISO-8601"\n}')
    doc.add_paragraph("Chuẩn hóa mã từ chối:")
    for text in ("no_face", "multiple_faces", "low_image_quality", "unknown_face", "ambiguous_match", "liveness_failed", "outside_attendance_window", "already_attended"):
        add_bullet(doc, text)

    doc.add_heading("2.3. Nghiệp vụ cần bổ sung", level=2)
    for text in (
        "Giáo viên xác nhận, thêm thủ công hoặc sửa trạng thái điểm danh.",
        "Mỗi lần sửa phải lưu người sửa, thời điểm, giá trị cũ/mới và lý do.",
        "Sinh viên xem lịch sử của chính mình và gửi yêu cầu xem xét.",
        "Xuất báo cáo theo lớp, buổi học, học kỳ và khoảng thời gian.",
        "Cấu hình múi giờ và quy tắc đi trễ thay vì giả định cứng.",
    ):
        add_bullet(doc, text)

    doc.add_heading("3. Tầng 2 - AI/ML correctness", level=1)
    doc.add_heading("3.1. Xác định đúng loại hệ thống", level=2)
    doc.add_paragraph(
        "Dự án không huấn luyện mô hình mới mà sử dụng mô hình pretrained của face_recognition/dlib. Do đó train/validation/test theo nghĩa huấn luyện không áp dụng trực tiếp. Tuy nhiên vẫn bắt buộc có enrollment, validation và test độc lập để hiệu chỉnh ngưỡng và đo chất lượng nhận diện."
    )
    add_table(
        doc,
        ["Thành phần", "Tính hợp lệ", "Rủi ro cần xử lý"],
        [
            ("Embedding 128 chiều", "Kỹ thuật chuẩn của thư viện", "Phụ thuộc mô hình pretrained và điều kiện ảnh"),
            ("Khoảng cách Euclidean", "Phù hợp baseline matching", "Không phải xác suất"),
            ("Top-1/Top-2 margin", "Giảm trường hợp mơ hồ", "Ngưỡng cần calibration"),
            ("EAR/chớp mắt", "Kiểm tra tương tác cơ bản", "Có thể bị video replay; không phải anti-spoof hoàn chỉnh"),
        ],
        [2100, 2940, 4320],
    )

    doc.add_heading("3.2. Thiết kế dữ liệu đánh giá", level=2)
    add_code(doc, "data/private/\n  enrollment/person_001/...\n  validation/known/...\n  validation/unknown/...\n  test/known/...\n  test/unknown/...\ndata/results/")
    for text in (
        "Không commit data/private, embedding hoặc database thật.",
        "Mỗi người có 3-5 ảnh enrollment; validation và test chụp ở thời điểm khác.",
        "Nhóm unknown phải gồm người chưa từng xuất hiện trong enrollment.",
        "Dùng SHA-256 để phát hiện file trùng giữa các tập.",
        "Ghi metadata: người, phiên chụp, ánh sáng, thiết bị và góc mặt; không lưu thông tin thừa.",
    ):
        add_bullet(doc, text)

    doc.add_heading("3.3. Ngăn data leakage", level=2)
    add_number(doc, "Tách theo phiên chụp hoặc theo người tùy mục tiêu đánh giá; không chia ngẫu nhiên các ảnh gần giống nhau.")
    add_number(doc, "Chọn threshold và margin chỉ trên validation.")
    add_number(doc, "Khóa cấu hình trước khi chạy test cuối.")
    add_number(doc, "Không xem test rồi quay lại điều chỉnh ngưỡng.")
    add_number(doc, "Lưu manifest và hash để kết quả có thể tái lập.")

    doc.add_heading("3.4. Đồng nhất preprocessing", level=2)
    doc.add_paragraph(
        "Enrollment và realtime cần dùng cùng chuyển đổi màu, detector, quy tắc số khuôn mặt, chuẩn hóa kích thước và tham số tạo embedding. Nếu realtime resize để tăng tốc, phải đánh giá chính pipeline đó thay vì đánh giá trên ảnh full-resolution rồi suy diễn."
    )
    add_code(doc, "def extract_embedding(image_rgb, config):\n    quality = validate_image_quality(image_rgb, config)\n    locations = detect_faces(image_rgb, config)\n    require_exactly_one_face(locations)\n    return encode_face(image_rgb, locations[0], config)")
    add_callout(doc, "Tiêu chí hoàn thành", "Một hàm/service dùng chung cho enrollment, offline evaluation và realtime inference; có test chứng minh cùng input tạo cùng embedding trong sai số cho phép.", GREEN)

    doc.add_heading("3.5. Calibration và metric", level=2)
    add_table(
        doc,
        ["Metric", "Công thức/ý nghĩa", "Mục đích"],
        [
            ("FAR", "Người lạ được chấp nhận / tổng lượt người lạ", "Đo nhận nhầm; nên ưu tiên thấp"),
            ("FRR", "Người hợp lệ bị từ chối / tổng lượt hợp lệ", "Đo bất tiện và số lần phải thử lại"),
            ("TAR", "Người hợp lệ được chấp nhận / tổng lượt hợp lệ", "Khả năng nhận đúng tại FAR mục tiêu"),
            ("Latency", "P50/P95/P99 theo frame hoặc request", "Đánh giá trải nghiệm và quy mô"),
            ("Ambiguous rate", "Top-1/Top-2 không đủ margin", "Đánh giá khả năng từ chối trường hợp mơ hồ"),
        ],
        [1500, 4320, 3540],
    )
    doc.add_paragraph("Quy trình chọn ngưỡng:")
    for text in (
        "Thử dải threshold, ví dụ 0.35-0.60, trên validation.",
        "Vẽ hoặc xuất bảng FAR/FRR theo threshold.",
        "Chọn operating point ưu tiên FAR thấp vì nhận nhầm người nghiêm trọng hơn yêu cầu thử lại.",
        "Chạy test đúng một lần với cấu hình đã khóa.",
        "Lưu JSON/CSV kết quả và phiên bản code/config đi kèm.",
    ):
        add_number(doc, text)

    doc.add_heading("3.6. Baseline và liveness", level=2)
    doc.add_paragraph("Bảng so sánh tối thiểu, chỉ điền bằng số liệu do script sinh:")
    add_table(
        doc,
        ["Phương pháp", "FAR", "FRR", "P95 latency"],
        [
            ("Distance threshold", "Chưa đo", "Chưa đo", "Chưa đo"),
            ("Threshold + margin", "Chưa đo", "Chưa đo", "Chưa đo"),
            ("Nhiều template/người", "Chưa đo", "Chưa đo", "Chưa đo"),
        ],
        [3960, 1800, 1800, 1800],
    )
    add_callout(doc, "Cách mô tả liveness", "Phát hiện chớp mắt là kiểm tra tương tác giúp giảm một số trường hợp dùng ảnh tĩnh; không thay thế Presentation Attack Detection chuyên dụng.", YELLOW)
    doc.add_paragraph("Nâng cấp theo thứ tự: challenge ngẫu nhiên → giới hạn thời gian → quay trái/phải → đánh giá replay → cân nhắc PAD chuyên dụng khi có dữ liệu phù hợp.")

    doc.add_heading("4. Tầng 3 - Software Engineering", level=1)
    doc.add_heading("4.1. Kiến trúc mục tiêu", level=2)
    add_code(doc, "UI / API\n   ↓\nApplication services\n   ↓\nDomain rules\n   ↓\nRepositories + face engine + database")
    doc.add_paragraph("Nguyên tắc phụ thuộc:")
    for text in (
        "UI không chứa SQL hoặc quyết định nhận diện.",
        "API chỉ validate, xác thực, gọi service và ánh xạ lỗi sang HTTP response.",
        "Recognition không trực tiếp ghi database.",
        "Một service nghiệp vụ dùng chung cho UI và API.",
        "Exception nghiệp vụ có kiểu rõ ràng thay vì chuỗi thông báo tự do.",
    ):
        add_bullet(doc, text)
    add_code(doc, "class StudentNotInRosterError(AttendanceError): ...\nclass DuplicateAttendanceError(AttendanceError): ...\nclass AttendanceWindowClosedError(AttendanceError): ...")

    doc.add_heading("4.2. Kiểm thử còn thiếu", level=2)
    add_table(
        doc,
        ["Nhóm", "Trường hợp bắt buộc", "Ưu tiên"],
        [
            ("Matcher", "127/129 chiều, NaN, Inf, rỗng, bằng threshold, bằng margin", "P0"),
            ("Ảnh", "file rỗng, sai MIME, không mặt, nhiều mặt, tối, mờ, ảnh rất lớn", "P0"),
            ("Điểm danh", "trùng, ngoài roster, ranh giới đi trễ, hai request đồng thời", "P0"),
            ("API", "thiếu/sai key, payload sai, ID không tồn tại, không lộ stack trace", "P0"),
            ("UI/WebRTC", "camera mất kết nối, frame lỗi, session restart", "P1"),
            ("Database", "migration, corrupted embedding, backup/restore, Unicode", "P1"),
        ],
        [1800, 5940, 1620],
    )

    doc.add_heading("4.3. Tooling và CI", level=2)
    for text in (
        "Ruff: lint và format.",
        "Mypy: kiểm tra type cho service, matcher và repository.",
        "pytest-cov: theo dõi coverage phần nghiệp vụ; mục tiêu ban đầu 70-80%.",
        "Bandit và pip-audit: phát hiện lỗi phổ biến và dependency có lỗ hổng.",
        "Lockfile: khóa dependency bằng uv/Poetry hoặc file requirements triển khai có version chính xác.",
    ):
        add_bullet(doc, text)
    add_callout(doc, "CI gate", "Syntax → Ruff → type-check → unit/integration tests → coverage → security scan → Docker build/smoke test.", GREEN)

    doc.add_heading("4.4. Cấu hình và xử lý lỗi", level=2)
    add_code(doc, "APP_ENV=development\nDATABASE_PATH=data/attendance.db\nAPI_KEY=change-me\nFACE_DISTANCE_THRESHOLD=0.50\nFACE_MARGIN_THRESHOLD=0.05\nBIOMETRIC_RETENTION_DAYS=365\nMAX_UPLOAD_SIZE_MB=10\nLOG_LEVEL=INFO")
    for text in (
        "Production phải từ chối khởi động nếu còn API key mặc định.",
        "Client nhận mã lỗi ổn định; chi tiết exception chỉ ghi log an toàn ở server.",
        "Không log embedding, PIN, API key hoặc dữ liệu ảnh.",
        "Dùng đường dẫn tương đối theo project/data, không dùng đường dẫn máy cá nhân.",
        "Thực hiện clean-clone test trong CI để chứng minh hướng dẫn cài đặt.",
    ):
        add_bullet(doc, text)

    doc.add_heading("5. Tầng 4 - Production và Business value", level=1)
    doc.add_heading("5.1. Security và privacy", level=2)
    add_table(
        doc,
        ["Kiểm soát", "Yêu cầu", "Mức ưu tiên"],
        [
            ("Authentication", "Tài khoản/role riêng; không dùng một API key chung", "P0"),
            ("Authorization", "Admin/Teacher/Student với phạm vi dữ liệu rõ", "P0"),
            ("Rate limit", "Login, enrollment, attendance và upload", "P0"),
            ("Transport", "HTTPS bắt buộc; secure headers", "P0"),
            ("Biometric", "Mã hóa, consent có phiên bản, retention và xóa", "P0"),
            ("Audit", "Log bất biến cho xem/sửa/xóa và thay đổi quyền", "P1"),
            ("Recovery", "Backup mã hóa và kiểm thử restore", "P1"),
        ],
        [1980, 5580, 1800],
    )
    doc.add_paragraph("Vai trò đề xuất:")
    for text in (
        "Admin: quản lý tài khoản, cấu hình và audit.",
        "Teacher: quản lý lớp/buổi học, xem và sửa điểm danh có lý do.",
        "Student: đăng ký hoặc xóa biometric của chính mình và xem lịch sử cá nhân.",
    ):
        add_bullet(doc, text)
    add_callout(doc, "Rủi ro API", "Nếu endpoint chấp nhận student_id và recognition_distance từ client có shared API key, người giữ key có thể tạo bản ghi giả. Quyết định nhận diện cần được thực hiện hoặc ký xác nhận ở phía server.", RED)

    doc.add_heading("5.2. Khả năng phục vụ 100 users", level=2)
    doc.add_paragraph(
        "Cần phân biệt 100 sinh viên được lưu trong database với 100 camera hoạt động đồng thời. Trường hợp đầu tương đối nhẹ; trường hợp sau chưa được chứng minh và có thể quá tải CPU, WebRTC và SQLite."
    )
    add_table(
        doc,
        ["Kịch bản", "Đo lường", "Điều kiện đạt đề xuất"],
        [
            ("10 phiên", "error rate, CPU/RAM, P95", "Không lỗi chức năng; P95 trong SLA demo"),
            ("50 phiên", "DB lock, queue, P95/P99", "Không mất bản ghi; degradation có kiểm soát"),
            ("100 phiên", "throughput, tài nguyên, WebRTC", "Chỉ công bố nếu test ổn định lặp lại"),
            ("Cùng sinh viên", "race và unique constraint", "Chỉ một bản ghi hợp lệ"),
        ],
        [1800, 3240, 4320],
    )
    doc.add_paragraph("Khi cần mở rộng:")
    for text in (
        "Chuyển SQLite sang PostgreSQL.",
        "Tách recognition worker/service khỏi Streamlit.",
        "Không xử lý mọi frame; giới hạn FPS và cache template có version.",
        "Dùng queue cho tác vụ nặng, nhiều API worker và reverse proxy.",
        "Theo dõi request ID, latency, error rate, CPU/RAM và database connection.",
    ):
        add_bullet(doc, text)

    doc.add_heading("5.3. Giá trị nghiệp vụ và demo", level=2)
    doc.add_paragraph("Demo CV nên chứng minh trọn vẹn một phiên sử dụng:")
    for text in (
        "Giáo viên tạo lớp và buổi học.",
        "Sinh viên đồng ý và đăng ký khuôn mặt.",
        "Sinh viên hợp lệ được nhận diện và ghi nhận.",
        "Người lạ hoặc kết quả mơ hồ bị từ chối.",
        "Ảnh tĩnh không hoàn thành challenge bị từ chối.",
        "Điểm danh lần hai không tạo bản ghi trùng.",
        "Giáo viên xem, sửa có lý do và xuất báo cáo.",
        "Sinh viên rút consent và embedding được xóa theo chính sách.",
    ):
        add_number(doc, text)

    doc.add_heading("6. Lộ trình triển khai theo sprint", level=1)
    add_table(
        doc,
        ["Sprint", "Mục tiêu", "Deliverable", "Definition of Done"],
        [
            ("1", "Correctness nền tảng", "Pipeline preprocessing chung; mã lỗi chuẩn; test P0", "Test xanh, không regression"),
            ("2", "Đánh giá AI", "Manifest dữ liệu, calibration/evaluation script", "Có FAR/FRR/TAR/latency thật"),
            ("3", "Clean engineering", "Service layer, exception, Ruff/mypy/coverage", "CI gate chạy trên clean clone"),
            ("4", "Demo portfolio", "GIF/video, sample an toàn, README kết quả", "Người lạ clone và chạy theo README"),
            ("5", "Security/business", "RBAC, rate limit, sửa điểm danh, audit", "Test quyền và luồng nghiệp vụ"),
            ("6", "Scale", "Load test, PostgreSQL/worker nếu cần", "Báo cáo P50/P95/P99 và giới hạn"),
        ],
        [1080, 2160, 3060, 3060],
    )

    doc.add_heading("7. Ma trận kiểm thử và tiêu chí nghiệm thu", level=1)
    add_table(
        doc,
        ["Hạng mục", "Bằng chứng bắt buộc", "Trạng thái hiện tại"],
        [
            ("Problem", "Use case, actor, luồng lỗi, manual override", "Đạt phần lớn"),
            ("No leakage", "Manifest/hash và quy tắc split", "Chưa có dữ liệu đánh giá"),
            ("Metric", "Script + CSV/JSON kết quả", "Chưa đạt"),
            ("Baseline", "Bảng so sánh có số thật", "Đạt một phần"),
            ("Edge cases", "Test matrix và CI log", "Đạt một phần"),
            ("Clean clone", "CI hoặc video chạy từ clone mới", "Đạt một phần"),
            ("Security", "RBAC/rate limit/privacy tests", "Đạt một phần"),
            ("100 concurrent", "Load-test report", "Chưa đạt"),
            ("Demo", "GIF/video và kịch bản tái lập", "Đạt một phần"),
            ("README", "Kiến trúc, run, metric, limitation", "Đạt"),
        ],
        [2160, 4680, 2520],
        status_col=2,
    )
    doc.add_paragraph("Checklist trước khi phát hành GitHub:")
    for text in (
        "[ ] Không có ảnh, embedding, database, secret hoặc đường dẫn cá nhân.",
        "[ ] pytest, lint, type-check và security scan đều xanh.",
        "[ ] Docker build và smoke test thành công từ clean clone.",
        "[ ] README không chứa metric chưa có bằng chứng.",
        "[ ] Demo dùng dữ liệu có consent hoặc dữ liệu minh họa an toàn.",
        "[ ] Limitations và intended use được trình bày rõ.",
        "[ ] License, .gitignore, .env.example và hướng dẫn xóa dữ liệu đầy đủ.",
    ):
        add_bullet(doc, text)

    doc.add_heading("8. Hướng dẫn trình bày trên GitHub và CV", level=1)
    doc.add_heading("8.1. README", level=2)
    for text in (
        "Problem statement và intended use.",
        "Ảnh/GIF demo và luồng end-to-end.",
        "Sơ đồ kiến trúc và trách nhiệm từng module.",
        "Quick start, Docker, API và cách chạy test.",
        "Thiết kế matching, unknown rejection và liveness.",
        "Evaluation protocol cùng metric thật.",
        "Security/privacy, limitations và roadmap.",
    ):
        add_bullet(doc, text)

    doc.add_heading("8.2. Cách viết trong CV", level=2)
    add_callout(
        doc,
        "Mô tả an toàn",
        "Xây dựng prototype điểm danh khuôn mặt bằng Python, Streamlit/WebRTC, FastAPI và SQLite; thiết kế pipeline embedding, unknown rejection bằng distance + margin, kiểm tra tương tác bằng chớp mắt, audit/retention và automated tests.",
        GREEN,
    )
    doc.add_paragraph("Chỉ thêm số liệu khi đã đo, ví dụ:")
    add_bullet(doc, "Đạt FAR/FRR/TAR trên bộ test độc lập gồm N người và M lượt - điền đúng số thực tế.")
    add_bullet(doc, "P95 latency bằng X ms trên cấu hình phần cứng Y - điền đúng báo cáo load test.")
    add_bullet(doc, "Coverage X% cho service/matcher/database - lấy trực tiếp từ CI.")
    add_callout(doc, "Không được ghi", "Accuracy 99%, production-ready, chống spoof an toàn hoặc hỗ trợ 100 users nếu chưa có bằng chứng tái lập.", RED)

    doc.add_heading("8.3. Thứ tự ưu tiên cuối cùng", level=2)
    priorities = (
        "Đồng nhất preprocessing và chuẩn hóa decision/error code.",
        "Viết evaluation/calibration script và tạo protocol chống leakage.",
        "Bổ sung edge-case/concurrency/API security tests.",
        "Tách nghiệp vụ khỏi UI/API và thêm quality gate trong CI.",
        "Hoàn thiện demo trực quan cùng README trung thực.",
        "Bổ sung RBAC, rate limit, bảo vệ biometric và manual correction.",
        "Load test trước khi thay database hoặc tuyên bố quy mô.",
        "Chỉ nâng cấp mô hình AI/PAD khi baseline và dữ liệu cho thấy cần thiết.",
    )
    for text in priorities:
        add_number(doc, text)

    doc.add_paragraph()
    add_callout(doc, "Kết luận", "Giá trị CV mạnh nhất đến từ bằng chứng có thể kiểm tra: benchmark thật, test đáng tin cậy, demo rõ ràng và README trung thực; không phải từ việc thêm nhiều thuật ngữ AI.", LIGHT_BLUE)

    props = doc.core_properties
    props.title = "Hướng dẫn cải thiện chi tiết Face Attendance System"
    props.subject = "Đánh giá và lộ trình cải thiện theo bốn tầng"
    props.author = "Face Attendance System"
    props.keywords = "face attendance, AI evaluation, software engineering, security, portfolio"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
